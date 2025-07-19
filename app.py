# app.py

import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO
from rules_engine import build_llm_prompt
import datetime

# --- PAGE CONFIGURATION ---
st.set_page_config(
    page_title="PoC NDA Generator (Gemini)",
    page_icon="✍️",
    layout="wide"
)

# --- API KEY & SESSION STATE ---
try:
    # Configure the Gemini client with the API key from secrets
    api_key=st.secrets["GEMINI_API_KEY"]
except (FileNotFoundError, KeyError):
    st.error("Secrets file not found or GEMINI_API_KEY is missing. Please create .streamlit/secrets.toml with your key.")
    st.stop()

if "nda_text" not in st.session_state:
    st.session_state.nda_text = ""
if "prompt" not in st.session_state:
    st.session_state.prompt = ""

# --- FUNCTIONS ---
def generate_contract_from_prompt(prompt):
    """Calls the Google Gemini API to generate the contract text."""
    
    try:
        # Configure the API key
        genai.configure(api_key=api_key)
        
        # Set up the model generation configuration
        generation_config = {
            "temperature": 0.3,
            "top_p": 1,
            "top_k": 1,
            "max_output_tokens": 8192,
        }

        # Initialize the Generative Model
        model = genai.GenerativeModel(
            model_name="gemini-2.5-flash-lite-preview-06-17",
            generation_config=generation_config
        )

        with st.spinner("Drafting the NDA..."):
            response = model.generate_content(prompt)
        
        return response.text
        
    except Exception as e:
        st.error(f"An error occurred with the Gemini API: {e}")
        return None

def create_docx(text):
    """Creates a Word document in memory from the given text."""
    doc = Document()
    doc.add_heading('Non-Disclosure Agreement', 0)
    # Simple parsing: add paragraphs based on newlines
    for para in text.split('\n'):
        # Check if the line is likely a heading (e.g., "Article 1. Definitions")
        if para.strip().lower().startswith(("article", "section")) or para.strip().endswith(":"):
            doc.add_heading(para.strip(), level=2)
        elif para.strip(): # Avoid adding empty paragraphs
            doc.add_paragraph(para)
    
    # Save the document to a byte stream
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

# --- UI LAYOUT ---
st.title("📄 Proof-of-Concept NDA Generator")
st.markdown("This tool generates a first draft of a Non-Disclosure Agreement based on your selections. **All generated content must be reviewed by qualified legal counsel.**")

st.header("Contract Details")
col1, spacer, col2 = st.columns([1, 0.1, 1])

with col1:
          st.header("First Party")
    
          first_party = st.text_input("First Party", "OCP")
          first_party_adress = st.text_input(
              "First Party Address",
              "Rue Al Abtal, Hay Erraha 20200, Casablanca Morocco"
          )
          first_party_incorporation_state = st.text_input(
              "First Party Incorporation State",
              "Casablanca, Morocco"
          )
          first_party_representative = st.text_input(
              "First Party Representative",
              "Mr. Last First"
          )
          first_party_registration_number = st.text_input(
              "First Party Registration Number",
              "40327"
          )
          first_party_role = st.selectbox(
              "First Party Role",
              ("Receiving Party", "Disclosing Party", "Both (Bilateral)")
          )

with col2:
          st.header("Second Party")
    
          second_party = st.text_input("Second Party", "UM6P Foundry")
          second_party_adress = st.text_input(
              "Second Party Address",
              "Lot 660, Hay Moulay Rachid Ben Guerir, 43150, Morocco"
          )
          second_party_incorporation_state = st.text_input(
              "Second Party Incorporation State",
              "Ben Guerir, Morocco"
          )
          second_party_representative = st.text_input(
              "Second Party Representative",
              "Mr. Last First"
          )
          second_party_registration_number = st.text_input(
              "Second Party Registration Number",
              "1037"
          )

          default_second_party_role = ""
          if first_party_role == "Disclosing Party":
              default_second_party_role = "Receiving Party"
          elif first_party_role == "Receiving Party":
              default_second_party_role = "Disclosing Party"
          elif first_party_role == "Both (Bilateral)":
              default_second_party_role = "Both (Bilateral)"
          second_party_role = st.text_input(
              "Second Party Role",
              value=default_second_party_role,
              disabled=True
          )
    
st.markdown("----------")
purpose_type = st.selectbox(
              "Purpose of Disclosure Type",
              ("Education", "Business", "Research", "Other")
          )        
default_purpose = ""
if purpose_type == "Research":
    default_purpose = "The contemplated cooperation between the Parties relating to the impact of AI on the future of work"
elif purpose_type == "Education":
    default_purpose = "The contemplated cooperation between the Parties relating to the creation of a joint Master on Data Science"
elif purpose_type == "Business":
    default_purpose = "The contemplated merger between the Parties"          
purpose = st.text_area(
            "Purpose of Disclosure",
            value=default_purpose,           
            height=100
          )        

applicable_law = st.selectbox(
              "Applicable Law",
              ("English Law", "French Law", "Moroccan Law")
          )

language = st.selectbox(
              "Language of the Contract",
              ("English", "French")
          )

duration = st.number_input(
              "Duration of Confidentiality (months)",
              min_value=1,
              max_value=60,
              value=36,
              help="Duration for which the confidentiality obligations will apply."
          )

date = st.date_input(
              "Effective Date",
              value=st.session_state.get("effective_date", datetime.date.today()),
              help="The date when the NDA becomes effective. Defaults to today."
          )
          
litigation = st.selectbox(
              "Dispute Resolution",
              ("Arbitration under ICC Rules, seat in Paris", "Arbitration under LCIA Rules, seat in London")
          )
          
submitted = st.button("Draft NDA", type="primary", use_container_width=True)


if submitted:
    user_inputs = {
        "first_party": first_party,
        "first_party_adress": first_party_adress,
        "first_party_incorporation_state": first_party_incorporation_state,
        "first_party_representative": first_party_representative,
        "first_party_registration_number": first_party_registration_number,
        "first_party_role": first_party_role,

        "second_party": second_party,
        "second_party_adress": second_party_adress,
        "second_party_incorporation_state": second_party_incorporation_state,
        "second_party_representative": second_party_representative,
        "second_party_registration_number": second_party_registration_number,
        "second_party_role": second_party_role,

        "purpose_type": purpose_type,
        "purpose": purpose,
        "applicable_law": applicable_law,
        "language": language,
        "duration": duration,
        "date": date.strftime("%Y-%m-%d"),
        "litigation": litigation,
        "nature_of_obligations": "Unilateral" if first_party_role != "Both (Bilateral)" else "Bilateral"
    }
    
    # Build the prompt (No changes needed in rules_engine.py)
    st.session_state.prompt = build_llm_prompt(user_inputs)
    
    # Generate the contract using the new Gemini function
    generated_text = generate_contract_from_prompt(st.session_state.prompt)
    if generated_text:
        st.session_state.nda_text = generated_text

st.markdown("----------")
st.header("Generated Document")
    
if st.session_state.nda_text:
    st.markdown(st.session_state.nda_text)

    col1, spacer, col2 = st.columns([0.9, 0.01, 0.09]) 
    with col1:
        feedback = st.text_input(
            "Enter your message:",
            placeholder="Type something here...",
            label_visibility="collapsed" 
        )

    with col2:
        chat_button = st.button("Send ✉️")

    if chat_button:
        if feedback:
            st.success(f"Message sent: '{feedback}'")
        else:
            st.warning("Please enter a message before sending!")

    # Create docx in memory for download
    docx_bytes = create_docx(st.session_state.nda_text)
        
    st.download_button(
        label="📥 Download as Word Document",
        data=docx_bytes,
        file_name=f"NDA_{first_party.replace(' ', '')}_{second_party.replace(' ', '')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
        
        # Expander to see the prompt that was used
    with st.expander("Show the AI Prompt"):
        st.code(st.session_state.prompt, language='markdown')
else:
    st.info("Fill out the form on the left and click 'Draft NDA' to generate the document.")